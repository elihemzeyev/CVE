import re
import requests
import threading
from flask import Flask, render_template, request, send_file
from selenium import webdriver
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from fpdf import FPDF
from docx import Document
import os
import webbrowser
# Ali Hamzayev was responsible for the end-to-end technical implementation of the project, ensuring that all system components were designed and developed to meet both functional and performance requirements.
app = Flask(__name__)

# Set up Selenium options for headless mode
chrome_options = webdriver.ChromeOptions()
chrome_options.add_argument("--disable-gpu")
chrome_options.add_argument("--no-sandbox")
chrome_options.add_argument("--headless")  # Enable headless mode

# Directory to save the generated files
DOWNLOAD_FOLDER = 'downloads'
os.makedirs(DOWNLOAD_FOLDER, exist_ok=True)

#Farida Dashdamirova wrote cve validation function which one of the important parts in project.
def validate_cve_id(cve_id):
    """Validate the format of the CVE ID."""
    pattern = r"CVE-\d{4}-\d{4,7}"
    return re.match(pattern, cve_id) is not None

#Aydin Khalilov fix MITRE api which we get CVE status also Vendor and Products.

def check_cve_status(cve_id):
    """Check if the CVE ID is published, reserved, or rejected."""
    url = f"https://cveawg.mitre.org/api/cve-id/{cve_id}"
    response = requests.get(url)
    if response.status_code == 200:
        data = response.json()
        return data.get('state')
    return None

def get_affected_assets(cve_id):
    """Get affected vendor and product information using the CVE API."""
    url = f"https://cveawg.mitre.org/api/cve/{cve_id}"
    response = requests.get(url)
    affected_assets = []
    
    if response.status_code == 200:
        data = response.json()
        try:
            for affected in data['containers']['cna']['affected']:
                vendor = affected['vendor']
                product = affected['product']
                affected_assets.append({"vendor": vendor, "product": product})
        except KeyError:
            print(f"Error parsing affected assets for CVE {cve_id}.")
    return affected_assets

def collect_data(cve_id, report):
    """Use Selenium to collect data from various sources."""
    service = Service(executable_path='chromedriver.exe')  # Update the path to your chromedriver
    driver = webdriver.Chrome(service=service, options=chrome_options)

    try:
        # Scrape NIST NVD for important information
        nvd_url = f"https://nvd.nist.gov/vuln/detail/{cve_id}"
        driver.get(nvd_url)

        try:
            # Use WebDriverWait to wait for elements to load
            cvss_panel_element = WebDriverWait(driver, 10).until(
                EC.presence_of_element_located((By.XPATH, "//*[@id='Vuln3CvssPanel']"))
            )
            
            # Locate CVSS score and vector using relative XPath
            cvss_score_element = cvss_panel_element.find_element(By.XPATH, ".//span[@class='severityDetail']")
            cvss_vector_element = cvss_panel_element.find_element(By.XPATH, ".//a[@id='Cvss3NistCalculatorAnchor']")
            
            # Extract and store CVSS score and vector
            report["cvss_score"] = cvss_score_element.text.strip()
            vector_full = cvss_vector_element.get_attribute("href").split('vector=')[1]
            report["cvss_details"] = vector_full.split('&')[0]
        
        except Exception as e:
            print(f"Error finding CVSS score and vector: {e}")

        try:
            # Wait for the description element to be visible using XPath
            description_element = WebDriverWait(driver, 10).until(
                EC.visibility_of_element_located((By.XPATH, "//p[@data-testid='vuln-description']"))
            )
            report["description"] = description_element.text.strip()
        except Exception as e:
            print(f"Error finding description: {e}")

        # Use the API to get affected assets
        report["affected_assets"] = get_affected_assets(cve_id)

        # Extract valid reference links
        report["references"] = get_valid_references(driver)
        # Use JavaScript to open the top 3 reference links in new tabs without focusing on them
        for link in report["references"][:3]:
            webbrowser.open_new_tab(link)
        
        # Scrape Exploit DB for exploits, including validation column
        exploitdb_url = f"https://www.exploit-db.com/search?cve={cve_id}"
        driver.get(exploitdb_url)
        try:
            WebDriverWait(driver, 10).until(
                EC.presence_of_all_elements_located((By.XPATH, "//*[@id='exploits-table']/tbody/tr"))
            )
            exploit_rows = driver.find_elements(By.XPATH, "//*[@id='exploits-table']/tbody/tr")
            for row in exploit_rows:
                title_element = row.find_element(By.XPATH, "./td[5]")
                title_text = title_element.text.strip()
                title_link_element = title_element.find_element(By.TAG_NAME, "a")
                title_link = title_link_element.get_attribute("href")
                download_link = title_link.replace("/exploits/", "/download/")  # Adjusted link replacement

                # Get validation status (e.g., green checkmark)
                validation_icon = row.find_element(By.XPATH, "./td[4]/i")
                validation = "Valid" if "mdi-check" in validation_icon.get_attribute("class") else "Invalid"

                report["exploits"].append({
                    "title": title_text, 
                    "title_link": title_link, 
                    "download_link": download_link,
                    "validation": validation
                })
        except Exception as e:
            print(f"No exploits found on Exploit DB for {cve_id}: {e}")

    finally:
        driver.quit()

def get_valid_references(driver):
    """Extracts up to 5 valid reference links from the NIST reference table."""
    valid_links = []
    try:
        reference_links = driver.find_elements(By.XPATH, "//table[contains(@data-testid, 'vuln-hyperlinks-table')]//a")
        for link in reference_links:
            if len(valid_links) >= 5:
                break
            href = link.get_attribute("href")
            if check_link_exists(href):
                valid_links.append(href)
    except Exception as e:
        print(f"Error finding references: {e}")
    return valid_links

def check_link_exists(url):
    """Checks if a URL exists by making an HTTP HEAD request."""
    try:
        response = requests.head(url, allow_redirects=True)
        return response.status_code == 200
    except requests.RequestException:
        return False

def start_data_collection(cve_id):
    """Initiate data collection in a separate thread."""
    report = {
        "cve_id": cve_id,
        "cvss_score": "",
        "cvss_details": "",
        "description": "",
        "affected_assets": [],
        "exploits": [],
        "references": []
    }
    thread = threading.Thread(target=collect_data, args=(cve_id, report))
    thread.start()
    thread.join()  # Wait for the thread to finish
    return report

#Ulvi Rzayev wrote generate_pdf function which we use to download reports in PDF format.

def generate_pdf(report):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)

    pdf.cell(200, 10, txt=f"CVE Report: {report['cve_id']}", ln=True, align='C')

    pdf.cell(200, 10, txt=f"CVSS Score: {report['cvss_score']}", ln=True)
    pdf.cell(200, 10, txt=f"CVSS Vector: {report['cvss_details']}", ln=True)
    pdf.multi_cell(0, 10, txt=f"Description: {report['description']}")

    pdf.cell(200, 10, txt="Affected Assets:", ln=True)
    for asset in report['affected_assets']:
        pdf.cell(200, 10, txt=f"Vendor: {asset['vendor']}, Product: {asset['product']}", ln=True)

    pdf.cell(200, 10, txt="Exploits:", ln=True)

    for exploit in report['exploits']:
        # Title
        pdf.set_font("Arial", style='B', size=12)  # Bold for titles
        pdf.multi_cell(0, 10, txt=f"Title: {exploit['title']}")
        
        # Reset font for the rest of the details
        pdf.set_font("Arial", size=12)
        
        # Link
        pdf.multi_cell(0, 10, txt=f"Link: {exploit['title_link']}")
        
        # Download link
        pdf.multi_cell(0, 10, txt=f"Download: {exploit['download_link']}")
        
        # Validation
        pdf.multi_cell(0, 10, txt=f"Validation: {exploit['validation']}")
        
        # Add some spacing between exploits
        pdf.ln(5)

    filename = os.path.join(DOWNLOAD_FOLDER, f"{report['cve_id']}.pdf")
    pdf.output(filename)
    return filename

def generate_docx(report):
    doc = Document()
    doc.add_heading(f"CVE Report: {report['cve_id']}", 0)

    doc.add_heading('CVSS Score', level=1)
    doc.add_paragraph(f"CVSS Score: {report['cvss_score']}")
    doc.add_paragraph(f"CVSS Vector: {report['cvss_details']}")

    doc.add_heading('Description', level=1)
    doc.add_paragraph(report['description'])

    doc.add_heading('Affected Assets', level=1)
    for asset in report['affected_assets']:
        doc.add_paragraph(f"Vendor: {asset['vendor']}, Product: {asset['product']}")

    doc.add_heading('Exploits', level=1)
    for exploit in report['exploits']:
        doc.add_paragraph(f"Title: {exploit['title']}, Link: {exploit['title_link']}, Download: {exploit['download_link']}, Validation: {exploit['validation']}")

    filename = os.path.join(DOWNLOAD_FOLDER, f"{report['cve_id']}.docx")
    doc.save(filename)
    return filename

def generate_html(report):
    html_content = render_template('report.html', report=report)
    filename = os.path.join(DOWNLOAD_FOLDER, f"{report['cve_id']}.html")
    with open(filename, 'w', encoding='utf-8') as f:
        f.write(html_content)
    return filename

@app.route('/', methods=['GET', 'POST'])
def index():
    if request.method == 'POST':
        cve_id = request.form['cve_id']
        if not validate_cve_id(cve_id):
            return render_template('index.html', error="Invalid CVE ID format. Please enter a valid CVE ID.")
        
        status = check_cve_status(cve_id)
        if status:
            if status in ['RESERVED', 'REJECTED']:
                return render_template('index.html', error=f"The CVE ID {cve_id} is {status}. Cannot proceed with data collection.")
            elif status != 'PUBLISHED':
                warning = "The CVE ID is not in a published state. Proceed with caution."
                report = start_data_collection(cve_id)
                
                # Save the generated files (PDF, DOCX, HTML) once during form submission
                generate_pdf(report)   # PDF file
                generate_docx(report)  # DOCX file
                generate_html(report)  # HTML file
                
                return render_template('report.html', report=report, warning=warning, status=status)
            
            report = start_data_collection(cve_id)
            
            # Save the generated files (PDF, DOCX, HTML) once during form submission
            generate_pdf(report)   # PDF file
            generate_docx(report)  # DOCX file
            generate_html(report)  # HTML file
            
            return render_template('report.html', report=report, status=status)
        else:
            return render_template('index.html', error=f"Could not determine the status of the CVE ID {cve_id}.")
    
    return render_template('index.html')


@app.route('/download/<filetype>/<cve_id>')
def download_file(filetype, cve_id):
    # Serve the already generated file based on the filetype
    if filetype == 'pdf':
        filename = os.path.join(DOWNLOAD_FOLDER, f"{cve_id}.pdf")
    elif filetype == 'docx':
        filename = os.path.join(DOWNLOAD_FOLDER, f"{cve_id}.docx")
    elif filetype == 'html':
        filename = os.path.join(DOWNLOAD_FOLDER, f"{cve_id}.html")
    else:
        return "Invalid file type requested.", 400

    # Ensure the file exists before sending
    if os.path.exists(filename):
        return send_file(filename, as_attachment=True)
    else:
        return "File not found.", 404




if __name__ == "__main__":
    app.run(debug=True)
